from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.shared import RGBColor



def emptyrowsloop(empty_rows: int, document: Document) -> Document:
    """loop for adding some empty rows"""
    for _ in range(empty_rows):
        document.add_paragraph()
    return document


def font_size(table: pd.DataFrame,font_size: int):
    """changes table's font size"""
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size= Pt(font_size)
    return 


def table_creation(input_table_columns: list,name_table_docx: list,table_int: pd.DataFrame,title: str, document: Document) -> Document:
    """Adding a table to the document with the title, and the set columns""" 
    overviewtransformationsHeader = document.add_heading(f'{title}', level=2)
    overviewtransformationsHeader.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document = emptyrowsloop(1, document)
    
    table = document.add_table(rows=len(table_int) + 1, cols=len(input_table_columns))
    table.style = 'Light Shading'
    
    for idx, name in enumerate(name_table_docx):
        table.cell(0, idx).text = name
    
    #adding data to the table
    for row_idx, (_, row) in enumerate(table_int[input_table_columns].iterrows(), start=1):
        for col_idx, value in enumerate(row):
            table.cell(row_idx, col_idx).text = str(value)
    # table formatting
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = 1
            cell.width = Inches(1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font_size(table, 10)
    document = emptyrowsloop(1, document)
    description = "<Placeholder for manual input>"
    paragraph = document.add_paragraph(description)
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
    document = emptyrowsloop(1, document)
    return document


def picture_creation(path: str, title: str, data_legend: dict, document: Document) -> Document:
    """Addition of a sankey picture with the accompanying title and legend explaining the nodes in the sankey diagram""" 
    overviewtransformationsHeader = document.add_heading(f'{title}', level=2)
    overviewtransformationsHeader.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document = emptyrowsloop(1, document)
    document.add_picture(path, width=Inches(7))
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    systemsummary = pd.DataFrame(data_legend)
    #Legend for systems
    document.add_heading('Legend:', level=3)

    systems_Legend = document.add_table(rows = len(systemsummary)*2, cols = 2)

    #adjust column width and row height
    systems_Legend.autofit = False
    systems_Legend.allow_autofit = False

    for cell in systems_Legend.columns[0].cells:
        cell.width = Inches(1)

    for cell in systems_Legend.columns[1].cells:
        cell.width = Inches(4)
        
    for row in systems_Legend.rows:
        row.height = Inches(0.18)
    row_idx = 0

    for i in range(len(systemsummary)*2):
        systems_Legend.rows[i].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    # Loop through the DataFrame rows
    Name = ['systemName']

    for _, row in systemsummary[Name].iterrows():
        # Increase the row index counter
        # Add data from input_table_columns to the table
        for col_idx, value in enumerate(row):
            systems_Legend.cell(row_idx, col_idx =1).text = ('- ', value)
        # skip one row
        row_idx += 2

    #Adding colors to the database
    for i in range(len(systemsummary)):
        #row 0
        cell_xml_element0 = systems_Legend.rows[0+2*i].cells[0]._tc
        table_cell_properties0 = cell_xml_element0.get_or_add_tcPr()
        shade_obj0 = OxmlElement('w:shd')
        shade_obj0.set(qn('w:fill'), systemsummary['Color'][i])
        table_cell_properties0.append(shade_obj0)

    font_size(systems_Legend, 10)
    return document


def front_page(document: Document, logopath: str, today: str) -> Document:
    """Creation of the first page with the logo, title and data"""
    document = emptyrowsloop(5, document)

    document.add_picture(logopath, width=Inches(2.5)) #logo

    #placing the logo in the center
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document = emptyrowsloop(3, document)

    document.add_heading('Rationalisation Model', 0)

    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document = emptyrowsloop(8, document)

    Datetext =  f'Date: {today}'
    Date = document.add_paragraph(Datetext)
    Date.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    document.add_page_break() #finishing the page/section 1
    return document

def heading(document: Document, title: str) -> Document:
    """ Adding the heading with the correct format"""
    summaryHeading = document.add_heading(title, level=1)
    summaryHeading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document = emptyrowsloop(1, document)
    return document

def main_report_generation(input_path:str, output_path:str):
    table11 = pd.read_csv(f'{input_path}/blocks_control.csv')
    table112 = pd.read_csv(f'{input_path}/blocks_dataflow.csv')
    table113 = pd.read_csv(f'{input_path}/source_df.csv')
    table2122 = pd.read_csv(f'{input_path}/target_df.csv')
    table211 = pd.read_csv(f'{input_path}/transformation_df.csv')
    table212 = pd.read_csv(f'{input_path}/split_df.csv')
    table22 = pd.read_csv(f'{input_path}/join_df.csv')
    logopath = f'{input_path}/MA logo.png'
    
    document = Document()
    document = front_page(document,logopath,str(datetime.now().date()))
    
    scope = 'SSIS'
    company = 'Rabobank'
    goal = 'The goal of this analysis is to uncover the complexity within the system in-scope and provide useful insights of its functionalities.'  
    
    #heading + intro 1
    document = heading(document, 'Summary of the analysis performed')
    document.add_paragraph(f'This document contains the results of the rationalisation model for the {company} with scope being {scope}. {goal}')
    document = emptyrowsloop(1, document)
    
    #table 1
    input_table_columns = ['FUNCTION', 'count']
    name_table_docx = ['Node task', 'Number of occurences']
    title = 'Overview of the nodes in the control flow'
    document = table_creation(input_table_columns, name_table_docx, table11, title, document)
    
    #sankey 1
    path_sankey = f'{input_path}/complete_flow.jpg'
    data_legend = {
        'systemName': ["Node", "External table", "Join or split node", "Filter node", "Variable", "Data transmission", "Transformation (existing column)", "Transformation (new column)"],
        'Color': ["#D0D3D3", "#42D6A4", "#9D94FF", "#DB59A5", "#D0D708", "#f0f8ff", "#FFB480", "#FF6961"]
    }
    title = "Complete lineage of the analysed SSIS package"
    picture_creation(path_sankey, title, data_legend, document)


    #sankey 2
    path_sankey = f'{input_path}/external_control.jpg'
    data_legend = {
        'systemName': ["External table", "Control flow node"],
        'Color': ["#42D6A4", "#D0D3D3"]
    }
    title = "External table connection to the control nodes"
    picture_creation(path_sankey, title, data_legend, document)
    
    #heading + intro 2
    document = heading(document, 'Details of the data flow Merge and filter')
    document.add_paragraph("This section zooms in on the critical observations derived from the analysis, with a specific emphasis on the data flow.")
    document = emptyrowsloop(1, document)
    
    #table 2
    input_table_columns = ['FUNCTION', 'count']
    name_table_docx = ['Node task', 'Number of occurences']
    title = 'Overview of the nodes in the data flow'
    document = table_creation(input_table_columns, name_table_docx, table112, title, document)
    
    #table 3
    input_table_columns = ['SOURCE_NAME', 'count']
    name_table_docx = ['Source table', 'Occurrences']
    title = 'Overview of utilised source tables in the data flow'
    document = table_creation(input_table_columns, name_table_docx, table113, title, document)
    
    #table 4
    input_table_columns = ['TARGET_NAME', 'count']
    name_table_docx = ['Target table', 'Occurrences']
    title = 'Overview of utilised target tables in the data flow'
    document = table_creation(input_table_columns, name_table_docx, table2122, title, document)
    
    #table 5
    input_table_columns = ['SOURCE_NAME','SOURCE_FIELD', 'TRANSFORMATION']
    name_table_docx = ['Node','Column', 'Transformation']
    title = 'Overview of transformations in the data flow'
    document = table_creation(input_table_columns, name_table_docx, table211, title, document)
    
    #table 6
    input_table_columns = ['LABEL_NODE','FUNCTION', 'SPLIT_ARG']
    name_table_docx = ['Node', 'Node task', 'Split argument']
    title = 'Overview of split arguments in the data flow'
    document = table_creation(input_table_columns, name_table_docx, table212, title, document)
    
    #table 7
    input_table_columns = ['LABEL_NODE','FUNCTION', 'JOIN_ARG']
    name_table_docx = ['Node', 'Node task', 'Join argument']
    title = 'Overview of join arguments in the data flow'
    document = table_creation(input_table_columns, name_table_docx, table22, title, document)
    
    #heading + intro 3
    document = heading(document, 'Sankey Diagrams')
    document.add_paragraph('This section contains the Merge and filter data flow in a sankey Diagram, giving you insights into the overall lineage and the transformations as well as model-identified focus points of the view.')
    
    #sankey 2
    path_sankey = f'{input_path}/sankey_dataflow.jpg'
    data_legend = {
        'systemName': ["Node", "External table", "Join or split node", "Variable", "Data transmission", "Transformation (existing column)", "Transformation (new column)"],
        'Color': ["#D0D3D3", "#42D6A4", "#9D94FF", "#D0D708", "#f0f8ff", "#FFB480", "#FF6961"]
    }
    title = "Lineage within the Merge and filter data flow"
    picture_creation(path_sankey, title, data_legend, document)
    document.save(output_path)
    return

if __name__ == "__main__":
    main_report_generation("C:/Users/Erwinsiegers/Documents/GitHub/ssis-code-parser/output-data/reports/tables", "C:/Users/Erwinsiegers/Documents/GitHub/ssis-code-parser/output-data/reports/MA_Rationalization_Model_Results.docx")
