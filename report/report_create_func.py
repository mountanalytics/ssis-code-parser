##Report for one system

####PACKAGES
from datetime import datetime
import docx as docx
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd #package for dataframes
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.shared import RGBColor
from docx.enum.section import WD_SECTION

############################################################ FUNCTIONS

#loop for adding some empty rows
def emptyrowsloop(empty_rows, document: Document) -> Document:
    for _ in range(empty_rows):
        document.add_paragraph()
    return document

# changes table's font size
def font_size(table,font_size):
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size= Pt(font_size)
    return 



def table_creation(SummaryColumns,column_name,table11,title, document) -> Document:
    OverviewTransformationsHeader = document.add_heading(f'{title}', level=2)
    OverviewTransformationsHeader.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document = emptyrowsloop(1, document)
    
    
    # TABLE Number of transformations per database
    
    
    #empty table for systems
    #SummaryColumns = ['Calculation view', 'Number of nodes', 'Number of transformations', 'Number of filters']
    Table = document.add_table(rows=len(table11) + 1, cols=len(SummaryColumns))
    Table.style = 'Light Shading'
    
    #adding column headers
    #column_name = ['Calculation view', 'Number of nodes', 'Number of transformations', 'Number of filters']
    for idx, name in enumerate(column_name):
        Table.cell(0, idx).text = name
    
    #adding data to the table
    for row_idx, (_, row) in enumerate(table11[SummaryColumns].iterrows(), start=1):
        for col_idx, value in enumerate(row):
            Table.cell(row_idx, col_idx).text = str(value)
    # Table formatting
    for row in Table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = 1
            cell.width = Inches(1)
    Table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    font_size(Table, 10)
    
    document = emptyrowsloop(1, document)
    
    ##################
    description = "<Placeholder for manual input>"
    paragraph = document.add_paragraph(description)
    
    # Add red color to the text
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
    
    document = emptyrowsloop(1, document)
    return document


def picture_creation(path: str, title: str, data: dict, document: Document) -> Document:
    OverviewTransformationsHeader = document.add_heading(f'{title}', level=2)
    OverviewTransformationsHeader.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document = emptyrowsloop(1, document)
    document.add_picture(path, width=Inches(7))
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    SystemSummary = pd.DataFrame(data)
    #Legend for systems
    document.add_heading('Legend:', level=3)

    Systems_Legend = document.add_table(rows = len(SystemSummary)*2, cols = 2)

    #adjust column width and row height
    Systems_Legend.autofit = False
    Systems_Legend.allow_autofit = False

    for cell in Systems_Legend.columns[0].cells:
        cell.width = Inches(1)

    for cell in Systems_Legend.columns[1].cells:
        cell.width = Inches(4)
        
    for row in Systems_Legend.rows:
        row.height = Inches(0.18)
    row_idx = 0

    for i in range(len(SystemSummary)*2):
        Systems_Legend.rows[i].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    # Loop through the DataFrame rows
    Name = ['SystemName']

    for _, row in SystemSummary[Name].iterrows():
        # Increase the row index counter
        # Add data from SummaryColumns to the table
        for col_idx, value in enumerate(row):
            Systems_Legend.cell(row_idx, col_idx =1).text = ('- ', value)
        # Skip one row
        row_idx += 2

    #Adding colors to the database
    for i in range(len(SystemSummary)):
        #row 0
        cell_xml_element0 = Systems_Legend.rows[0+2*i].cells[0]._tc
        table_cell_properties0 = cell_xml_element0.get_or_add_tcPr()
        shade_obj0 = OxmlElement('w:shd')
        shade_obj0.set(qn('w:fill'), SystemSummary['Color'][i])
        table_cell_properties0.append(shade_obj0)

    font_size(Systems_Legend, 10)
    return document










############################################################ BEGINNING OF THE DOCUMENT - Section 1

 #create document
def front_page(document: Document, logopath: str, today: str) -> Document:
    document = emptyrowsloop(5, document)

    document.add_picture(logopath, width=Inches(2.5)) #logo

    #placing the logo in the center
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document = emptyrowsloop(3, document)

    document.add_heading('Rationalisation Model', 0)

    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document = emptyrowsloop(8, document)

    DateText =  f'Date: {today}'
    Date = document.add_paragraph(DateText)
    Date.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    document.add_page_break() #finishing the page/section 1
    return document

def main_report_generation():
    path = "report/data"
    table11 = pd.read_csv(f'{path}/blocks_control.csv')
    table112 = pd.read_csv(f'{path}/blocks_dataflow.csv')
    table113 = pd.read_csv(f'{path}/source_df.csv')
    table2122 = pd.read_csv(f'{path}/target_df.csv')
    table211 = pd.read_csv(f'{path}/transformation_df.csv')
    table212 = pd.read_csv(f'{path}/split_df.csv')
    table22 = pd.read_csv(f'{path}/join_df.csv')
    
    # Mount Analytics logo
    logopath = f'{path}/MA logo.png'
    
    
    ###### Manual Text input
    
    # Scope
    Scope = 'SSIS'
    
    # Company name
    Company = 'Rabobank'
    
    # Todays date
    today = str(datetime.now().date())
    
    # The goal of this exercise (write a proper sentence here with a dot)
    Goal = 'The goal of this analysis is to uncover the complexity within the system in-scope and provide useful insights of its functionalities.'
    
    
    
    document = Document()
    document = front_page(document,logopath,today)
    # Heading of this section
    SummaryHeading = document.add_heading('Summary of the analysis performed', level=1)
    SummaryHeading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document = emptyrowsloop(1, document)
    
    # Summary paragraph
    document.add_paragraph(f'This document contains the results of the rationalisation model for the {Company} with scope being {Scope}. {Goal}')
    document = emptyrowsloop(1, document)
    
    
    SummaryColumns = ['FUNCTION', 'count']
    #adding column headers
    column_name = ['Node task', 'Number of occurences']
    title = 'Overview of the nodes in the control flow'
    document = table_creation(SummaryColumns, column_name, table11, title, document)
    
    
    path_sankey = f'{path}/external_control.jpg'
    data = {
        'SystemName': ["External table", "Control flow node"],
        'Color': ['black', 'aliceblue']
    }
    title = "External table connection to the control nodes"
    picture_creation(path_sankey, title, data, document)
    
    SummaryHeading = document.add_heading('Details of the data flow Merge and filter', level=1)
    SummaryHeading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document = emptyrowsloop(1, document)
    # Summary paragraph
    document.add_paragraph("This section zooms in on the critical observations derived from the analysis, with a specific emphasis on the data flow.")
    document = emptyrowsloop(1, document)
    
    SummaryColumns = ['FUNCTION', 'count']
    #adding column headers
    column_name = ['Node task', 'Number of occurences']
    title = 'Overview of the nodes in the data flow'
    document = table_creation(SummaryColumns, column_name, table112, title, document)
    
    SummaryColumns = ['SOURCE_NAME', 'count']
    #adding column headers
    column_name = ['Source table', 'Occurrences']
    title = 'Overview of utilised source tables in the data flow'
    document = table_creation(SummaryColumns, column_name, table113, title, document)
    
    
    
    SummaryColumns = ['TARGET_NAME', 'count']
    #adding column headers
    column_name = ['Target table', 'Occurrences']
    title = 'Overview of utilised target tables in the data flow'
    document = table_creation(SummaryColumns, column_name, table2122, title, document)
    
    
    SummaryColumns = ['SOURCE_NAME','SOURCE_FIELD', 'TRANSFORMATION']
    #adding column headers
    column_name = ['Node','Column', 'Transformation']
    title = 'Overview of transformations in the data flow'
    document = table_creation(SummaryColumns, column_name, table211, title, document)
    
    
    SummaryColumns = ['LABEL_NODE','FUNCTION', 'SPLIT_ARG']
    #adding column headers
    column_name = ['Node', 'Node task', 'Split argument']
    title = 'Overview of split arguments in the data flow'
    document = table_creation(SummaryColumns, column_name, table212, title, document)
    
    SummaryColumns = ['LABEL_NODE','FUNCTION', 'JOIN_ARG']
    #adding column headers
    column_name = ['Node', 'Node task', 'Join argument']
    title = 'Overview of join arguments in the data flow'
    document = table_creation(SummaryColumns, column_name, table22, title, document)
    
    
    
    document.add_heading('Sankey Diagrams', level=1)
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    document.add_paragraph('This section contains the Merge and filter data flow in a Sankey Diagram, giving you insights into the overall lineage and the transformations as well as model-identified focus points of the view.')
    path_sankey = f'{path}/sankey_dataflow.jpg'
    data = {
        'SystemName': ["Node", "External table", "join or split node", "Variable", "Data transmission", "Transformation (existing column)", "Transformation (new column)"],
        'Color': ['black', 'gold', 'dodgerblue', 'green', 'aliceblue', 'orangered', 'darkred']
    }
    title = "Lineage within the Merge and filter data flow"
    picture_creation(path_sankey, title, data, document)
    document.save("MA_Rationalization_Model_Results.docx")
    return

if __name__ == "__main__":
    main_report_generation()
