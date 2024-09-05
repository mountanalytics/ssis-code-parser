##Report for one system

####PACKAGES
from datetime import datetime
import docx as docx
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd #package for dataframes
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.shared import RGBColor
from docx.enum.section import WD_SECTION

############################################################ FUNCTIONS

#loop for adding some empty rows
def emptyrowsloop(empty_rows):
    for _ in range(empty_rows):
        document.add_paragraph()
    return

# changes table's font size
def font_size(table,font_size):
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size= Pt(font_size)
    return 



def table_creation(SummaryColumns,column_name,table11,title):
    OverviewTransformationsHeader = document.add_heading(f'{title}', level=2)
    OverviewTransformationsHeader.alignment = WD_ALIGN_PARAGRAPH.CENTER
    emptyrowsloop(1)
    
    
    # TABLE Number of transformations per database
    
    
    #empty table for systems
    #SummaryColumns = ['Calculation view', 'Number of nodes', 'Number of transformations', 'Number of filters']
    Table = document.add_table(rows=len(table11) + 1, cols=len(SummaryColumns))
    Table.style = 'Light Shading'
    
    #adding column headers
    #column_name = ['Calculation view', 'Number of nodes', 'Number of transformations', 'Number of filters']
    for idx, name in enumerate(column_name):
        Table.cell(0, idx).text = name
    
    #adding data to the table
    for row_idx, (_, row) in enumerate(table11[SummaryColumns].iterrows(), start=1):
        for col_idx, value in enumerate(row):
            Table.cell(row_idx, col_idx).text = str(value)
    # Table formatting
    for row in Table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = 1
            cell.width = Inches(1)
    Table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    font_size(Table, 10)
    
    emptyrowsloop(1)
    
    ##################
    description = "<Placeholder for manual input>"
    paragraph = document.add_paragraph(description)
    
    # Add red color to the text
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
    
    emptyrowsloop(1)
    return










table11 = pd.read_csv('data/blocks_control.csv')
table112 = pd.read_csv('data/blocks_dataflow.csv')
table113 = pd.read_csv('data/source_df.csv')
table2122 = pd.read_csv('data/target_df.csv')
table211 = pd.read_csv('data/transformation_df.csv')
table212 = pd.read_csv('data/split_df.csv')
table22 = pd.read_csv('data/join_df.csv')

# Mount Analytics logo
LogoPath = 'data/MA logo.png'


###### Manual Text input

# Scope
Scope = 'SSIS'

# Company name
Company = 'Rabobank'

# Todays date
Today = str(datetime.now().date())

# The goal of this exercise (write a proper sentence here with a dot)
Goal = 'The goal of this analysis is to uncover the complexity within the system in-scope and provide useful insights of its functionalities.'


############################################################ BEGINNING OF THE DOCUMENT - Section 1

document = Document() #create document

emptyrowsloop(5)

document.add_picture(LogoPath, width=Inches(2.5)) #logo

#placing the logo in the center
last_paragraph = document.paragraphs[-1] 
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
emptyrowsloop(3)

document.add_heading('Rationalisation Model', 0)

last_paragraph = document.paragraphs[-1] 
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

emptyrowsloop(8)

DateText =  f'Date: {Today}'
Date = document.add_paragraph(DateText)
Date.alignment = WD_ALIGN_PARAGRAPH.RIGHT

document.add_page_break() #finishing the page/section 1

############################################################ SUMMARY OF THE ANALYSIS - Section 2

# Heading of this section
SummaryHeading = document.add_heading('Summary of the analysis performed', level=1)
SummaryHeading.alignment = WD_ALIGN_PARAGRAPH.CENTER
emptyrowsloop(1)

# Summary paragraph
Summary = document.add_paragraph(f'This document contains the results of the rationalisation model for the {Company} with scope being {Scope}. {Goal}')
emptyrowsloop(1)


SummaryColumns = ['FUNCTION', 'count']
#adding column headers
column_name = ['Node task', 'Number of occurences']
title = 'Overview of the nodes in the control flow'
table_creation(SummaryColumns, column_name, table11, title)

SummaryHeading = document.add_heading('Details of the data flow Merge and filter', level=1)
SummaryHeading.alignment = WD_ALIGN_PARAGRAPH.CENTER
emptyrowsloop(1)
# Summary paragraph
Summary = document.add_paragraph("This section zooms in on the critical observations derived from the analysis, with a specific emphasis on the data flow.")
emptyrowsloop(1)

SummaryColumns = ['FUNCTION', 'count']
#adding column headers
column_name = ['Node task', 'Number of occurences']
title = 'Overview of the nodes in the data flow'
table_creation(SummaryColumns, column_name, table112, title)




SummaryColumns = ['SOURCE_NAME', 'count']
#adding column headers
column_name = ['Source table', 'Occurrences']
title = 'Overview of utilised source tables in the data flow'
table_creation(SummaryColumns, column_name, table113, title)



SummaryColumns = ['TARGET_NAME', 'count']
#adding column headers
column_name = ['Target table', 'Occurrences']
title = 'Overview of utilised target tables in the data flow'
table_creation(SummaryColumns, column_name, table2122, title)


SummaryColumns = ['SOURCE_NAME','SOURCE_FIELD', 'TRANSFORMATION']
#adding column headers
column_name = ['Node','Column', 'Transformation']
title = 'Overview of transformations in the data flow'
table_creation(SummaryColumns, column_name, table211, title)


SummaryColumns = ['LABEL_NODE','FUNCTION', 'SPLIT_ARG']
#adding column headers
column_name = ['Node', 'Node task', 'Split argument']
title = 'Overview of split arguments in the data flow'
table_creation(SummaryColumns, column_name, table212, title)

SummaryColumns = ['LABEL_NODE','FUNCTION', 'JOIN_ARG']
#adding column headers
column_name = ['Node', 'Node task', 'Join argument']
title = 'Overview of join arguments in the data flow'
table_creation(SummaryColumns, column_name, table22, title)

#######################################################
"""
SummaryGraph = document.add_picture("data/data_calc_view.jpg", width=Inches(7))
last_paragraph = document.paragraphs[-1] 
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

data = {
    'SystemName': ['Data source', 'Calculation view'],
    'Color': ['black', 'skyblue']
}

SystemSummary = pd.DataFrame(data)
#Legend for systems
document.add_heading('Legend:', level=3)

Systems_Legend = document.add_table(rows = len(SystemSummary)*2, cols = 2)

#adjust column width and row height
Systems_Legend.autofit = False
Systems_Legend.allow_autofit = False

for cell in Systems_Legend.columns[0].cells:
    cell.width = Inches(1)

for cell in Systems_Legend.columns[1].cells:
    cell.width = Inches(2)
    
for row in Systems_Legend.rows:
    row.height = Inches(0.18)
row_idx = 0

for i in range(len(SystemSummary)*2):
    Systems_Legend.rows[i].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

data = {
    'SystemName': ['Data source', 'Calculation view'],
    'Color': ['black', 'skyblue']
}

SystemSummary = pd.DataFrame(data)

# Loop through the DataFrame rows
Name = ['SystemName']

for _, row in SystemSummary[Name].iterrows():
    # Increase the row index counter
    # Add data from SummaryColumns to the table
    for col_idx, value in enumerate(row):
        Systems_Legend.cell(row_idx, col_idx =1).text = ('- ', value)
    # Skip one row
    row_idx += 2

#Adding colors to the database
for i in range(len(SystemSummary)):
    #row 0
    cell_xml_element0 = Systems_Legend.rows[0+2*i].cells[0]._tc
    table_cell_properties0 = cell_xml_element0.get_or_add_tcPr()
    shade_obj0 = OxmlElement('w:shd')
    shade_obj0.set(qn('w:fill'), SystemSummary['Color'][i])
    table_cell_properties0.append(shade_obj0)

font_size(Systems_Legend, 10)
"""
#######################################################
"""




SummaryColumns = ['Index','Column', "Value"]
#adding column headers
column_name = ['Calculation view 1','Calculation view 2', 'Similarity']
title = 'Similarity between analysed calculation views'
table_creation(SummaryColumns, column_name, table23, title)



SummaryHeading = document.add_heading('Detailed analysis of the Q_AccountsPayable calculation view', level=1)
SummaryHeading.alignment = WD_ALIGN_PARAGRAPH.CENTER
emptyrowsloop(1)

# Summary paragraph
Summary = document.add_paragraph("In this section, we dissect the Q_AccountsPayable calculation view, examining its components in detail. We highlight and analyze all the transformations and data sources incorporated within this calculation view. Furthermore, we provide a snapshot of the complete technical lineage, showcasing joins, filters, and transformations to offer a comprehensive understanding of its functionality and structure.")
emptyrowsloop(1)




table31 = table31[table31['Calculation view'] == 'Q_AccountsPayable']
SummaryColumns = ['Calculation view','Data source','Columns used','Columns in source','Percentage columns used']
#adding column headers
column_name = ['Calculation view','Data source','Columns used','Columns in source','Percentage columns used']
title = 'Detailed data source usage Q_AccountsPayable calculation view'
table_creation(SummaryColumns, column_name, table31, title)




SummaryColumns = ['Transformation','Occurrences']
#adding column headers
column_name = ['Transformation','Occurrences']
title = 'Transformations in the Q_AccountsPayable calculation view'
table_creation(SummaryColumns, column_name, substring_df, title)


SummaryColumns = ['SOURCE_NODE', 'SOURCE_FIELD', 'TARGET_NODE', 'TARGET_FIELD', 'TRANSFORMATION', 'JOIN_ARGU', 'FILTER', 'Complexity_Score']
#adding column headers
column_name = ['Source node', 'Source column', 'Target node', 'Target column', 'transformation', 'Join arguments', 'Filter', 'Complexity score']
title = 'Summarised technical lineage for the Q_AccountsPayable calculation view'
table_creation(SummaryColumns, column_name, Data_final_tech, title)


"""



############################################################ FOURTH SECTION - Sankey diagrams

document.add_heading('Sankey Diagrams', level=1)
last_paragraph = document.paragraphs[-1] 
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

document.add_paragraph('This section contains the Merge and filter data flow in a Sankey Diagram, giving you insights into the overall lineage and the transformations as well as model-identified focus points of the view.')

SummaryGraph = document.add_picture("data/sankey_dataflow.jpg", width=Inches(7))
last_paragraph = document.paragraphs[-1] 
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

data = {
    'SystemName': ["Node", "External table", "join or split node", "Variable", "Data transmission", "Transformation (existing column)", "Transformation (new column)"],
    'Color': ['black', 'gold', 'dodgerblue', 'green', 'aliceblue', 'orangred', 'darkred']
}

SystemSummary = pd.DataFrame(data)
#Legend for systems
document.add_heading('Legend:', level=3)

Systems_Legend = document.add_table(rows = len(SystemSummary)*2, cols = 2)

#adjust column width and row height
Systems_Legend.autofit = False
Systems_Legend.allow_autofit = False

for cell in Systems_Legend.columns[0].cells:
    cell.width = Inches(1)

for cell in Systems_Legend.columns[1].cells:
    cell.width = Inches(4)
    
for row in Systems_Legend.rows:
    row.height = Inches(0.18)
row_idx = 0

for i in range(len(SystemSummary)*2):
    Systems_Legend.rows[i].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY


data = {
    'SystemName': ["Node", "External table", "join or split node", "Variable", "Data transmission", "Transformation (existing column)", "Transformation (new column)"],
    'Color': ['black', 'gold', 'dodgerblue', 'green', 'aliceblue', 'orangered', 'darkred']
}

SystemSummary = pd.DataFrame(data)

# Loop through the DataFrame rows
Name = ['SystemName']

for _, row in SystemSummary[Name].iterrows():
    # Increase the row index counter
    # Add data from SummaryColumns to the table
    for col_idx, value in enumerate(row):
        Systems_Legend.cell(row_idx, col_idx =1).text = ('- ', value)
    # Skip one row
    row_idx += 2

#Adding colors to the database
for i in range(len(SystemSummary)):
    #row 0
    cell_xml_element0 = Systems_Legend.rows[0+2*i].cells[0]._tc
    table_cell_properties0 = cell_xml_element0.get_or_add_tcPr()
    shade_obj0 = OxmlElement('w:shd')
    shade_obj0.set(qn('w:fill'), SystemSummary['Color'][i])
    table_cell_properties0.append(shade_obj0)

font_size(Systems_Legend, 10)

document.save("MA_Rationalization_Model_Results.docx")